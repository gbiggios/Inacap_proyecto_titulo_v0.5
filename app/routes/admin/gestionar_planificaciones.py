from io import BytesIO
from docx import Document
from flask import Blueprint, render_template, request, redirect, send_file, url_for, flash
from ...extensions import db
from ...models import Planificacion, Taller
from flask_login import current_user, login_required
from app.routes.admin.decorators import admin_required
from app.forms import PlanificacionForm
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# Crear el Blueprint para las rutas de Planificación
planificacion_bp = Blueprint('planificacion_admin', __name__)

# Ruta para listar todas las planificaciones (solo admin)
@planificacion_bp.route('/', methods=['GET'])
@login_required
def listar_planificaciones_docente():
    # Obtener los talleres asignados al docente actual
    talleres = Taller.query.filter_by(id_docente=current_user.id_docente).all()

    if not talleres:
        flash('No tienes talleres asignados para mostrar planificaciones.')
        return redirect(url_for('docente.dashboard'))  # Redirigir al Dashboard si no hay talleres

    # Seleccionar el primer taller como predeterminado
    taller_seleccionado = talleres[0]

    # Obtener las planificaciones del taller seleccionado
    planificaciones = Planificacion.query.filter_by(taller_id=taller_seleccionado.id).all()

    # Listado de meses
    meses = ["marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]

    # Diccionario para almacenar las planificaciones por mes
    planificaciones_por_mes = {mes: None for mes in meses}

    # Llenar el diccionario con las planificaciones existentes
    for plan in planificaciones:
        planificaciones_por_mes[plan.mes] = plan

    # Determinar si se puede crear una nueva planificación
    siguiente_mes = None
    puede_crear = True
    for i, mes in enumerate(meses):
        if planificaciones_por_mes[mes] is None:
            siguiente_mes = mes
            if i > 0 and planificaciones_por_mes[meses[i - 1]] is None:
                puede_crear = False
            break

    return render_template(
        'user/gestionar_planificaciones_docente.html',
        talleres=talleres,
        taller_seleccionado=taller_seleccionado,
        planificaciones_por_mes=planificaciones_por_mes,  # Pasar las planificaciones por mes
        siguiente_mes=siguiente_mes,
        puede_crear=puede_crear
    )


# Ruta para seleccionar un taller específico
@planificacion_bp.route('/<int:taller_id>', methods=['GET'], endpoint='listar_planificaciones_taller')
@login_required
@admin_required
def listar_planificaciones_taller(taller_id):
    taller_seleccionado = Taller.query.get_or_404(taller_id)
    talleres = Taller.query.all()
    form= PlanificacionForm()
    planificaciones = {
        planificacion.mes: planificacion for planificacion in Planificacion.query.filter_by(taller_id=taller_seleccionado.taller_id).all()
    }

    meses = ["marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    siguiente_mes = None
    puede_crear = True

    for i, mes in enumerate(meses):
        if mes not in planificaciones:
            siguiente_mes = mes
            if i > 0 and meses[i - 1] not in planificaciones:
                puede_crear = False
            break

    return render_template(
        'admin/gestionar_planificaciones.html',
        talleres=talleres,
        taller_seleccionado=taller_seleccionado,
        planificaciones=planificaciones,
        siguiente_mes=siguiente_mes,
        puede_crear=puede_crear, form=form
    )

# Ruta para crear una nueva planificación (solo admin)
@planificacion_bp.route('/nueva', methods=['POST'], endpoint='crear_planificacion_admin')
@login_required
@admin_required
def crear_planificacion_admin():
    taller_id = request.form['taller_id']
    form= PlanificacionForm()
    
    # Validar que se haya seleccionado un taller válido
    if not taller_id or not taller_id.isdigit() or not Taller.query.get(taller_id):
        flash('Por favor, selecciona un taller válido.')
        return redirect(url_for('admin.planificacion_admin.listar_planificaciones_admin'))

    nueva_planificacion = Planificacion(
        taller_id=taller_id,
        mes=request.form['mes'],
        habilidades=','.join(request.form['habilidades'].split(',')),  # Convertir lista a cadena
        recursos=','.join(request.form['recursos'].split(',')),  # Convertir lista a cadena
        actividades=','.join(request.form['actividades'].split(',')),  # Convertir lista a cadena
        estado=request.form['estado']
    )
    db.session.add(nueva_planificacion)
    db.session.commit()
    flash('Planificación creada exitosamente')

    return redirect(url_for('admin.planificacion_admin.listar_planificaciones_taller', taller_id=taller_id,form=form))

# Ruta para editar una planificación existente (solo admin)
@planificacion_bp.route('/editar', methods=['POST'], endpoint='editar_planificacion_admin')
@login_required
@admin_required
def editar_planificacion_admin():
    planificacion = Planificacion.query.get_or_404(request.form['id'])

    form= PlanificacionForm()

    taller_id = request.form['taller_id']
    
    # Validar que se haya seleccionado un taller válido
    if not taller_id or not taller_id.isdigit() or not Taller.query.get(taller_id):
        flash('Por favor, selecciona un taller válido.')
        return redirect(url_for('admin.planificacion_admin.listar_planificaciones_taller', taller_id=planificacion.taller_id))

    planificacion.taller_id = taller_id
    planificacion.mes = request.form['mes']
    planificacion.habilidades = ','.join(request.form['habilidades'].split(','))  # Convertir lista a cadena
    planificacion.recursos = ','.join(request.form['recursos'].split(','))  # Convertir lista a cadena
    planificacion.actividades = ','.join(request.form['actividades'].split(','))  # Convertir lista a cadena
    planificacion.estado = request.form['estado']
    db.session.commit()
    flash('Planificación actualizada exitosamente')

    return redirect(url_for('admin.planificacion_admin.listar_planificaciones_taller', taller_id=taller_id, form=form))

# Ruta para eliminar una planificación (solo admin)
@planificacion_bp.route('/eliminar', methods=['POST'], endpoint='eliminar_planificacion_admin')
@login_required
@admin_required
def eliminar_planificacion_admin():
    planificacion = Planificacion.query.get_or_404(request.form['id'])
    taller_id = planificacion.taller_id
    db.session.delete(planificacion)
    db.session.commit()
    flash('Planificación eliminada exitosamente')

    return redirect(url_for('admin.planificacion_admin.listar_planificaciones_taller', taller_id=taller_id))



# Ruta para actualizar el objetivo general de un taller (solo admin)
@planificacion_bp.route('/actualizar_objetivo_general', methods=['POST'], endpoint='actualizar_objetivo_general')
@login_required
@admin_required
def actualizar_objetivo_general():
    taller_id = request.form['taller_id']
    taller = Taller.query.get_or_404(taller_id)

    objetivo_general = request.form['objetivo_general']
    taller.objetivo_general = objetivo_general

    db.session.commit()
    flash('Objetivo general actualizado exitosamente.')

    return redirect(url_for('admin.planificacion_admin.listar_planificaciones_taller', taller_id=taller_id))
from io import BytesIO
from flask import send_file
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

@planificacion_bp.route('/exportar/<int:taller_id>', methods=['GET'], endpoint='exportar_planificacion')
@login_required
@admin_required
def exportar_planificacion(taller_id):
    taller = Taller.query.get_or_404(taller_id)
    
    # Obtener el docente responsable del taller
    docente = taller.docente  # Asumiendo que la relación está configurada en el modelo Taller

    planificaciones = Planificacion.query.filter_by(taller_id=taller_id).order_by(Planificacion.mes).all()

    # Crear el documento de Word
    doc = Document()
    
    # Encabezado principal
    heading = doc.add_heading('PLANIFICACIÓN ANUAL ACLE', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Información de taller y responsable
    info_paragraph = doc.add_paragraph()
    info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = info_paragraph.add_run(f"TALLER: {taller.nombre}\n")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'

    # Extraer el nombre completo del docente
    nombre_docente = f"{docente.nombre} {docente.apellido_paterno}" if docente else "N/A"
    run = info_paragraph.add_run(f"RESPONSABLE: {nombre_docente}\n")
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run = info_paragraph.add_run(f"OBJETIVO GENERAL: {taller.objetivo_general}\n\n")
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    # Crear tabla para las planificaciones mensuales
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Mes'
    hdr_cells[1].text = 'Habilidades'
    hdr_cells[2].text = 'Recursos'
    hdr_cells[3].text = 'Actividades'

    # Formato de encabezado de tabla
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(12)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Meses y contenido de planificación
    meses = ["marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    for mes in meses:
        planificacion_mes = next((p for p in planificaciones if p.mes == mes), None)
        row_cells = table.add_row().cells
        row_cells[0].text = mes.capitalize()

        if planificacion_mes:
            row_cells[1].text = planificacion_mes.habilidades.replace(",", ", ")
            row_cells[2].text = planificacion_mes.recursos.replace(",", ", ")
            row_cells[3].text = planificacion_mes.actividades.replace(",", ", ")
        else:
            row_cells[1].text = 'N/A'
            row_cells[2].text = 'N/A'
            row_cells[3].text = 'N/A'

        # Formato de celdas
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.size = Pt(10)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Guardar en memoria y exportar
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name=f"Planificacion_{taller.nombre}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

